import os, sys, shutil, tempfile
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import docx
from caiguard import core, store, docx_io

def test_roundtrip():
    tmp = tempfile.mkdtemp()
    path = os.path.join(tmp, "agreement.docx")
    d = docx.Document(); d.add_heading("Agreement", 1)
    d.add_paragraph("The Contractor MUST maintain $2,000,000 of coverage.")
    d.add_paragraph("Payment SHALL be made within 30 days.")
    d.save(path)

    m = core.enroll(path); assert len(m["sections"]) == 3

    # external Word edit (weaken MUST -> SHOULD)
    dd = docx.Document(path)
    for p in dd.paragraphs:
        if "MUST maintain" in p.text:
            p.runs[0].text = p.text.replace("MUST", "SHOULD")
            for r in p.runs[1:]: r.text = ""
    dd.save(path)
    _, pend = core.verify(path)
    assert any(x["level"] == "control-weakened" for x in pend), pend

    # program edit is written into the .docx
    sid = [s["id"] for s in store.load(path)["sections"] if "Payment" in s["baseline"]][0]
    core.apply_edit(path, sid, "Payment SHALL be made within 10 days.")
    assert any("10 days" in p.text for p in docx.Document(path).paragraphs)

    # reject restores the .docx
    _, pend2 = core.verify(path)
    pid = [p["id"] for p in pend2 if p["actor"] == "program"][0]
    core.reject(path, pid)
    assert any("30 days" in p.text for p in docx.Document(path).paragraphs)
    shutil.rmtree(tmp)
    print("roundtrip: OK")

if __name__ == "__main__":
    test_roundtrip()
