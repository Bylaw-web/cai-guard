"""Read sections from a .docx (body + headers + footers) and write edits back into it."""
import os, re, time

def is_docx(path): return path.lower().endswith(".docx")

# OneDrive/Word artifacts that are NOT real user edits of the enrolled document
_ARTIFACT = re.compile(r"-(?:DESKTOP|LAPTOP|PC)-[A-Z0-9]{4,}\.|conflicted copy|\(\d+\)\.docx$", re.I)

def is_ignorable(path):
    """True for temp/lock/conflict files that should never be treated as an edit event."""
    b = os.path.basename(path)
    if b.startswith("~$") or b.startswith(".~caiguard-"):   # Word lock file / our own atomic temp
        return True
    return bool(_ARTIFACT.search(b))

def is_hydrated(path):
    """False only if this is a OneDrive 'cloud-only' placeholder not downloaded locally.

    Reading a placeholder would force a (possibly slow) download, so callers can skip it.
    Best-effort and Windows-specific; anything uncertain returns True.
    """
    try:
        import ctypes
        OFFLINE = 0x1000
        RECALL_ON_DATA_ACCESS = 0x400000
        attrs = ctypes.windll.kernel32.GetFileAttributesW(str(path))
        if attrs == -1:
            return True
        return not (attrs & (OFFLINE | RECALL_ON_DATA_ACCESS))
    except Exception:
        return True

def _locked_msg(path):
    onedrive = "onedrive" in path.lower()
    tail = " If the file is in OneDrive, pause syncing or wait for it to finish, then try again." if onedrive else ""
    return ("The document is locked — it's open in Microsoft Word or another program. "
            "Close it there, then try again." + tail)

def _save_docx(doc, path, attempts=3):
    """Save atomically, with a short retry so transient OneDrive/AV locks don't surface as a crash.

    Writes to a temp file in the same folder, then os.replace()s it into place — so a crash or a
    sync interruption can never leave a half-written .docx, and the visible-lock window is minimal.
    """
    d = os.path.dirname(os.path.abspath(path))
    tmp = os.path.join(d, f".~caiguard-{os.getpid()}-{os.path.basename(path)}")
    last = None
    for i in range(attempts):
        try:
            doc.save(tmp)
            os.replace(tmp, path)      # atomic on the same volume
            return True
        except PermissionError as e:
            last = e
            time.sleep(0.4 * (i + 1))
        finally:
            if os.path.exists(tmp):
                try: os.remove(tmp)
                except OSError: pass
    raise RuntimeError(_locked_msg(path)) from last

def _iter_docx(doc):
    """Yield (id, region, paragraph, style)."""
    for i, p in enumerate(doc.paragraphs):
        yield (f"b{i}", "body", p, (p.style.name if p.style else "") or "")
    for si, sec in enumerate(doc.sections):
        for i, p in enumerate(sec.header.paragraphs):
            yield (f"h{si}_{i}", "header", p, "header")
        for i, p in enumerate(sec.footer.paragraphs):
            yield (f"f{si}_{i}", "footer", p, "footer")

def read_sections(path):
    if is_docx(path):
        import docx
        doc = docx.Document(path)
        out = []
        for sid, region, p, style in _iter_docx(doc):
            if not p.text.strip():
                continue
            heading = style.startswith("Heading") or style == "Title"
            out.append({"id": sid, "region": region, "heading": heading,
                        "style": style, "text": p.text})
        return out
    # markdown / plain text: blocks separated by blank lines
    raw = open(path, encoding="utf-8").read().split("\n")
    blocks, cur = [], []
    for ln in raw:
        if ln.strip() == "":
            if cur: blocks.append("\n".join(cur)); cur = []
        else:
            cur.append(ln)
    if cur: blocks.append("\n".join(cur))
    return [{"id": f"b{i}", "region": "body", "heading": b.lstrip().startswith("#"),
             "style": "md", "text": b} for i, b in enumerate(blocks) if b.strip()]

def _locate(doc, sid):
    if sid.startswith("b") and sid[1:].isdigit():
        return doc.paragraphs[int(sid[1:])]
    m = re.match(r"([hf])(\d+)_(\d+)", sid)
    si, i = int(m.group(2)), int(m.group(3))
    sec = doc.sections[si]
    return (sec.header if m.group(1) == "h" else sec.footer).paragraphs[i]

def write_section(path, section_id, new_text):
    """Write new_text into a body/header/footer paragraph and save."""
    if is_docx(path):
        import docx
        from docx.oxml.ns import qn
        try:
            doc = docx.Document(path)
        except PermissionError:
            raise RuntimeError(_locked_msg(path))
        para = _locate(doc, section_id)
        # Replace ALL textual content (runs, hyperlinks, bookmarks) but keep paragraph properties.
        # p.text includes hyperlink text while p.runs does not, so the old runs[0]-only rewrite
        # corrupted any paragraph containing a hyperlink. Clearing every child except w:pPr fixes it.
        p = para._p
        for child in list(p):
            if child.tag == qn("w:pPr"):
                continue
            p.remove(child)
        para.add_run(new_text)
        _save_docx(doc, path)
        return True
    # markdown block write-back
    raw = open(path, encoding="utf-8").read().split("\n")
    blocks, cur, order = [], [], []
    for ln in raw:
        if ln.strip() == "":
            if cur: blocks.append("\n".join(cur)); cur = []
        else:
            cur.append(ln)
    if cur: blocks.append("\n".join(cur))
    idx = int(section_id[1:])
    if idx < len(blocks): blocks[idx] = new_text
    open(path, "w", encoding="utf-8").write("\n\n".join(blocks))
    return True
